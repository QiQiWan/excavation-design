from __future__ import annotations

from pathlib import Path

from docx import Document

from app.reports.docx_report import export_docx_report
from app.schemas.domain import (
    CalculationResult,
    GoverningValues,
    Point2D,
    Polyline2D,
    Project,
    StabilityDetailedResult,
)
from app.services.core_engineering_presentation import (
    build_core_standard_guidance,
    build_scheme_comparison,
    build_stability_distribution,
)
from app.services.core_workspace import build_core_workspace_status
from app.services.delivery_package import _write_manager_summary
from app.services.design_service import auto_diaphragm_wall
from app.services.excavation_service import make_excavation_model


def _project() -> Project:
    excavation = make_excavation_model(
        "visibility",
        Polyline2D(
            points=[
                Point2D(x=0, y=0), Point2D(x=48, y=0),
                Point2D(x=48, y=22), Point2D(x=0, y=22),
            ],
            closed=True,
        ),
        0.0,
        -11.0,
    )
    return Project(name="V3.41可视化工程", excavation=excavation, retainingSystem=auto_diaphragm_wall(excavation))


def _calculated_project() -> Project:
    project = _project()
    project.calculation_results.append(
        CalculationResult(
            projectId=project.id,
            caseId="CASE-1",
            governingValues=GoverningValues(
                maxDisplacement=18.6,
                maxSupportAxialForce=2450.0,
                maxWallMoment=860.0,
                embedmentSafetyFactorMin=1.42,
                heaveSafetyFactorMin=1.31,
                seepageSafetyFactorMin=1.24,
            ),
            stabilityDetailedResult=StabilityDetailedResult(
                heaveFactor=1.31,
                seepageFactor=1.24,
                confinedUpliftFactor=1.18,
                overallStabilityFactor=1.36,
                minSafetyFactor=1.18,
                controllingMode="confined_uplift",
            ),
            checks=[
                {
                    "ruleId": "STAB-UPLIFT",
                    "category": "承压水突涌稳定",
                    "calculatedValue": {"value": 1.18, "source": "layered-water"},
                    "limitValue": 1.10,
                    "standard": "JGJ 120-2012",
                    "message": "承压水突涌安全系数",
                }
            ],
            checkSummary={"fail": 0, "warning": 1, "pass": 8},
        )
    )
    return project


def test_core_standard_guidance_is_compact_and_stage_aligned() -> None:
    guidance = build_core_standard_guidance()
    assert set(guidance) == {"basis", "input", "scheme", "calculation", "reinforcement", "deliverables"}
    assert all(len(rows) <= 5 for rows in guidance.values())
    assert any("JGJ 120" in str(row.get("code")) for row in guidance["calculation"])


def test_stability_distribution_exposes_factor_limit_margin_and_standard() -> None:
    result = build_stability_distribution(_calculated_project())
    assert result["summary"]["count"] >= 4
    assert result["summary"]["minimumFactor"] == 1.18
    uplift = next(item for item in result["factors"] if item["code"] == "uplift")
    assert uplift["limit"] == 1.1
    assert uplift["marginRatio"] > 1.0
    assert "JGJ 120" in uplift["standard"]


def test_core_status_includes_visualization_and_scheme_inputs() -> None:
    status = build_core_workspace_status(_calculated_project(), {"workspaceBytes": 4096})
    assert status["standards"]["calculation"]
    assert status["stabilityDistribution"]["summary"]["minimumFactor"] == 1.18
    assert "rows" in status["schemeComparison"]


def test_manager_readable_docx_uses_plain_language_sections(tmp_path: Path) -> None:
    path = export_docx_report(_calculated_project(), tmp_path)
    document = Document(path)
    text = "\n".join(p.text for p in document.paragraphs)
    text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "项目管理摘要" in text
    assert "多方案比选" in text
    assert "稳定性安全系数" in text
    assert "设计流程—关键计算—规范条文对应矩阵" in text
    assert "JSON" not in text[: max(0, text.find("附录"))]


def test_manager_summary_prioritizes_actions_and_hides_machine_data(tmp_path: Path) -> None:
    project = _calculated_project()
    for folder in ("00_release", "30_reports", "10_drawings/quick_review", "40_rebar"):
        (tmp_path / folder).mkdir(parents=True, exist_ok=True)
    path = _write_manager_summary(
        tmp_path,
        project,
        [
            {"category": "计算", "check": "工程计算", "status": "warning", "action": "复核控制工况"},
            {"category": "图纸", "check": "图纸完整性", "status": "pass", "action": "—"},
        ],
        "review_complete",
    )
    content = path.read_text(encoding="utf-8")
    assert "项目管理摘要" in content
    assert "稳定性安全系数" in content
    assert "复核控制工况" in content
    assert "机器数据集中在 50_data 和 90_audit" in content


def test_scheme_comparison_has_consistent_empty_state() -> None:
    result = build_scheme_comparison(_project())
    assert result["candidateCount"] == 0
    assert result["comparisonAvailable"] is False


def test_diverse_scheme_search_returns_structural_alternatives() -> None:
    from app.services.support_layout_optimizer import optimize_support_layout_candidates

    _, candidates = optimize_support_layout_candidates(
        _project(),
        max_candidates=3,
        search_config={"requireDiverseSchemes": True, "maxTrials": 48},
    )
    assert len(candidates) >= 2
    signatures = {
        (
            item.variable_summary.get("topologyFamily"),
            item.column_count,
            item.support_count,
            item.target_spacing,
        )
        for item in candidates
    }
    assert len(signatures) >= 2
    assert all(item.variable_summary.get("diversityBasis") for item in candidates)
