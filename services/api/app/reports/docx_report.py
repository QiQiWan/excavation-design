from __future__ import annotations

from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.reports.charts import generate_report_charts
from app.schemas.domain import Project
from app.services.core_engineering_presentation import (
    build_core_standard_guidance,
    build_scheme_comparison,
    build_stability_distribution,
    build_verification_distribution,
)
from app.services.design_basis import build_design_basis
from app.services.standards_matrix import build_standards_process_matrix
from app.version import SOFTWARE_VERSION, version_manifest


DISCLAIMER = (
    "本计算书由软件依据项目输入、当前规则库和已实现的计算模型自动生成，用于方案设计、"
    "技术复核和项目沟通。正式施工图、专家论证和报审前，必须由具备相应资质的岩土及结构"
    "专业人员复核输入资料、计算模型、规范适用性、施工阶段和构造措施。"
)

STATUS_TEXT = {
    "pass": "通过",
    "warning": "预警",
    "fail": "不通过",
    "manual_review": "需人工复核",
    "missing_input": "缺资料",
    "not_calculated": "待重算",
    "not_applicable": "不适用",
    "not_implemented": "未实现",
    "missing": "缺失",
    "ready": "已完成",
    "completed": "已完成",
}


def _text(value: Any, fallback: str = "-") -> str:
    if value is None:
        return fallback
    if isinstance(value, float):
        return f"{value:,.3f}".rstrip("0").rstrip(".")
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, dict):
        for key in ("value", "calculatedValue", "result", "factor", "ratio"):
            if key in value:
                return _text(value.get(key), fallback)
        return fallback
    return str(value)


def _status(value: Any) -> str:
    key = str(value or "").strip().lower()
    return STATUS_TEXT.get(key, str(value or "未判定"))


def _number(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, dict):
        for key in ("value", "calculatedValue", "result", "factor", "ratio"):
            if key in value:
                return _number(value.get(key))
        return None
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    return number


def _short(value: Any, limit: int = 90) -> str:
    value_text = _text(value).replace("\n", " ")
    return value_text if len(value_text) <= limit else value_text[: limit - 1] + "…"

def _human_rule_name(item: dict[str, Any]) -> str:
    rule_id = str(item.get("ruleId") or item.get("category") or "检查")
    exact = {
        "GB50010-NODE-BEARING-SUBSET": "支撑端部及围檩节点局部承压",
        "JGJ120-2012-BASE-HEAVE-SUBSET": "坑底抗隆起稳定",
        "JGJ120-2012-SEEPAGE-STABILITY-SUBSET": "渗流稳定",
        "JGJ120-2012-CONFINED-WATER-UPLIFT-SCREEN": "承压水突涌稳定",
        "JGJ120-2012-OVERALL-STABILITY-CIRCULAR-SCREEN": "整体稳定",
        "JGJ120-2012-4.2-EMBEDMENT-STABILITY-SCREEN": "围护墙嵌固稳定",
    }
    if rule_id in exact:
        return exact[rule_id]
    token_map = (
        ("WALE-DEFLECTION", "围檩挠度"),
        ("WALE-NODE-REBAR", "围檩节点附加钢筋"),
        ("WALE-SHEAR", "围檩抗剪"),
        ("WALE-FLEXURE", "围檩抗弯"),
        ("WALE_SUPPORT_BAY", "围檩支点间距"),
        ("SUPPORT-DEEP-DESIGN-STABILITY", "支撑构件强度与稳定"),
        ("SUPPORT-CONSTRUCTION-EFFECTS", "支撑预加轴力及施工效应"),
        ("LAYOUT-SCREEN-SPAN", "支撑跨度与构造"),
        ("GEOLOGICAL-EXTRAPOLATION", "地质模型外推范围"),
        ("LAYOUT-SCREEN-SPACING", "支撑间距"),
    )
    upper = rule_id.upper()
    for token, label in token_map:
        if token in upper:
            return label
    return rule_id


def _human_standard_reference(item: dict[str, Any]) -> str:
    rule_id = str(item.get("ruleId") or "")
    raw = str(item.get("clauseReference") or item.get("clause_reference") or item.get("standardReference") or "").strip()
    upper = rule_id.upper()
    if rule_id == "GB50010-NODE-BEARING-SUBSET" or "local compression/detailing subset" in raw.lower():
        return "GB/T 50010-2010（2024局部修订）：局部受压、节点附加钢筋和构造要求；正式条文适用性需复核"
    if "WALE-DEFLECTION" in upper:
        return "围檩正常使用极限状态：挠度限值按项目控制标准和适用规范确认"
    if "WALE-NODE-REBAR" in upper:
        return "GB/T 50010-2010（2024局部修订）：围檩节点锚固、附加钢筋与构造协调"
    if "WALE-SHEAR" in upper:
        return "GB/T 50010-2010（2024局部修订）：钢筋混凝土围檩抗剪与箍筋构造"
    if "WALE-FLEXURE" in upper:
        return "GB/T 50010-2010（2024局部修订）：钢筋混凝土围檩正截面受弯与配筋构造"
    if "SUPPORT-DEEP-DESIGN-STABILITY" in upper:
        return "JGJ 120-2012、GB 50017-2017、GB/T 50010-2010：支撑传力、构件稳定和节点构造"
    if "WALE_SUPPORT_BAY" in upper:
        return "围檩支点间距项目硬限值及支撑传力连续性控制"
    if "SUPPORT-CONSTRUCTION-EFFECTS" in upper:
        return "JGJ 120-2012：内支撑预加轴力、温度效应和施工偏差控制"
    if "LAYOUT-SCREEN-SPAN" in upper or "LAYOUT-SCREEN-SPACING" in upper:
        return "JGJ 120-2012 第4.7节：内支撑布置、跨度、间距和构造原则"
    if "GEOLOGICAL-EXTRAPOLATION" in upper:
        return "GB 55017-2021：工程勘察资料完整性、外推边界和补充勘察要求"
    if rule_id.startswith("JGJ120") and not raw:
        return "JGJ 120-2012：基坑支护结构计算与稳定性验算"
    return _short(raw or "见规则编号和规范对应矩阵", 90)


def _group_review_checks(checks: list[dict[str, Any]], limit: int = 16) -> list[list[Any]]:
    groups: dict[tuple[str, str, str], dict[str, Any]] = {}
    for item in checks:
        status = str(item.get("status") or "manual_review")
        rule_id = str(item.get("ruleId") or item.get("category") or "检查")
        message = _short(item.get("message"), 130)
        key = (status, rule_id, message)
        group = groups.setdefault(key, {
            "status": status,
            "item": item,
            "count": 0,
            "values": [],
            "limits": [],
            "objects": set(),
        })
        group["count"] += 1
        value = _number(item.get("calculatedValue") if "calculatedValue" in item else item.get("calculated_value"))
        threshold = _number(item.get("limitValue") if "limitValue" in item else item.get("limit_value"))
        if value is not None:
            group["values"].append(value)
        if threshold is not None:
            group["limits"].append(threshold)
        object_id = item.get("objectId") or item.get("object_id")
        if object_id:
            group["objects"].add(str(object_id))

    order = {"fail": 0, "warning": 1, "manual_review": 2, "pass": 3}
    ranked = sorted(groups.values(), key=lambda row: (order.get(row["status"], 4), -row["count"], _human_rule_name(row["item"])))
    rows: list[list[Any]] = []
    for group in ranked[:limit]:
        item = group["item"]
        values = list(group["values"])
        thresholds = list(group["limits"])
        haystack = f"{item.get('ruleId') or ''} {item.get('message') or ''}".lower()
        lower_controls = any(token in haystack for token in ("安全系数", "safety factor", "稳定", "stability", "抗隆起", "嵌固"))
        governing = min(values) if values and lower_controls else max(values) if values else None
        threshold = max(thresholds) if thresholds and lower_controls else min(thresholds) if thresholds else None
        count_text = f"共 {group['count']} 条"
        if group["objects"]:
            count_text += f"，涉及 {len(group['objects'])} 个对象"
        message = f"{_short(item.get('message'), 105)}；{count_text}。"
        rows.append([
            group["status"],
            _human_rule_name(item),
            f"{_text(governing)} / {_text(threshold)}",
            message,
            _human_standard_reference(item),
        ])
    return rows


def _set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_cell_text(cell, value: Any, *, bold: bool = False, size: float = 8.5, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(_text(value))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(
    document: Document,
    headers: list[str],
    rows: Iterable[Iterable[Any]],
    *,
    widths: list[float] | None = None,
    font_size: float = 8.2,
    status_column: int | None = None,
) -> None:
    row_values = [list(row) for row in rows]
    if not row_values:
        row_values = [["-"] * len(headers)]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0]
    _set_repeat_table_header(header)
    for index, label in enumerate(headers):
        _set_cell_text(header.cells[index], label, bold=True, size=font_size, color="FFFFFF")
        _set_cell_fill(header.cells[index], "355B83")
        if widths and index < len(widths):
            header.cells[index].width = Inches(widths[index])
    for row in row_values:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            text = _status(value) if status_column == index else value
            _set_cell_text(cells[index], text, size=font_size)
            if status_column == index:
                normalized = str(value or "").lower()
                if normalized == "fail":
                    _set_cell_fill(cells[index], "FDE2E2")
                elif normalized == "warning":
                    _set_cell_fill(cells[index], "FFF1CC")
                elif normalized == "pass":
                    _set_cell_fill(cells[index], "E4F5E8")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_key_value_table(document: Document, rows: list[tuple[str, Any]], columns: int = 2) -> None:
    table = document.add_table(rows=0, cols=columns * 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for offset in range(0, len(rows), columns):
        cells = table.add_row().cells
        group = rows[offset: offset + columns]
        for index in range(columns):
            label_cell = cells[index * 2]
            value_cell = cells[index * 2 + 1]
            if index < len(group):
                label, value = group[index]
                _set_cell_text(label_cell, label, bold=True, size=8.6)
                _set_cell_fill(label_cell, "EAF0F6")
                _set_cell_text(value_cell, value, size=9)
            else:
                _set_cell_text(label_cell, "", size=8.6)
                _set_cell_text(value_cell, "", size=9)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_picture(document: Document, chart: dict[str, str] | None, caption: str, width: float = 6.2) -> None:
    if not chart:
        return
    try:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(chart["path"], width=Inches(width))
        caption_p = document.add_paragraph(caption)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.style = document.styles["Caption"] if "Caption" in document.styles else document.styles["Normal"]
    except Exception as exc:
        document.add_paragraph(f"图表未能插入：{exc}")


def _flatten_checks(project: Project) -> list[dict[str, Any]]:
    if not project.calculation_results:
        return []
    latest = project.calculation_results[-1]
    output: list[dict[str, Any]] = []
    for item in latest.checks or []:
        output.append(item if isinstance(item, dict) else item.model_dump(mode="json", by_alias=True))
    for stage in latest.stage_results or []:
        for item in stage.checks or []:
            output.append(item if isinstance(item, dict) else item.model_dump(mode="json", by_alias=True))
    unique: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str, str]] = set()
    for item in output:
        key = (
            str(item.get("ruleId") or item.get("rule_id") or ""),
            str(item.get("objectId") or item.get("object_id") or ""),
            str(item.get("status") or ""),
            _text(item.get("calculatedValue") if "calculatedValue" in item else item.get("calculated_value")),
        )
        if key not in seen:
            seen.add(key)
            unique.append(item)
    return unique


def _governing_conclusion(project: Project, checks: list[dict[str, Any]]) -> str:
    latest = project.calculation_results[-1] if project.calculation_results else None
    if latest is None:
        return "当前尚未形成完整计算结果，暂不能给出工程通过性结论。"
    fail_count = int((latest.check_summary or {}).get("fail", 0) or 0)
    warning_count = int((latest.check_summary or {}).get("warning", 0) or 0)
    if fail_count:
        return f"当前方案存在 {fail_count} 个不通过项，需调整结构体系、截面或施工条件后重新计算。"
    if warning_count:
        return f"当前计算未发现不通过项，但存在 {warning_count} 个预警项，应完成针对性复核后再进入正式出图。"
    if not checks:
        return "当前没有可追溯的规范检查记录，计算结果只能作为方案参考。"
    return "当前已实现的规范检查未发现不通过项，可进入配筋和审查成果编制；正式发行仍需专业校审。"


def _styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (
        ("Title", 24, "17324D"),
        ("Heading 1", 16, "17324D"),
        ("Heading 2", 13, "2E5578"),
        ("Heading 3", 11, "355B83"),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    if "Caption" in styles:
        styles["Caption"].font.name = "Arial"
        styles["Caption"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles["Caption"].font.size = Pt(8.5)
        styles["Caption"].font.color.rgb = RGBColor(90, 105, 120)


def _header_footer(document: Document, project: Project) -> None:
    for section in document.sections:
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        header = section.header.paragraphs[0]
        header.text = f"{project.name} - 基坑围护结构计算书"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 116, 139)
        footer = section.footer.paragraphs[0]
        footer.text = f"PitGuard V{SOFTWARE_VERSION} | 方案设计与技术复核辅助"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 116, 139)


def export_docx_report(project: Project, output_dir: str | Path) -> Path:
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    path = output / f"{project.id}_calculation_report.docx"
    document = Document()
    _styles(document)
    _header_footer(document, project)

    latest = project.calculation_results[-1] if project.calculation_results else None
    checks = _flatten_checks(project)
    check_counts = Counter(str(item.get("status") or "unknown") for item in checks)
    standards = build_core_standard_guidance()
    scheme_comparison = build_scheme_comparison(project)
    stability = build_stability_distribution(project)
    verification = build_verification_distribution(project)
    design_basis = build_design_basis(project)
    formal_scenarios = dict((project.advanced_engineering or {}).get("formalAdverseScenarioSuite") or {})
    p3_closure = dict((project.advanced_engineering or {}).get("p3DetailingClosure") or {})
    try:
        charts = generate_report_charts(project, output)
    except Exception:
        charts = []
    chart_by_title = {item.get("title"): item for item in charts}

    # Cover page
    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("基坑围护结构设计计算书")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("方案设计 · 施工阶段计算 · 稳定性分析 · 配筋深化").italic = True
    document.add_paragraph()
    _add_key_value_table(document, [
        ("项目名称", project.name),
        ("项目地点", project.location or "未录入"),
        ("软件版本", SOFTWARE_VERSION),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("基坑深度", f"{_text(getattr(project.excavation, 'depth', None))} m" if project.excavation else "未录入"),
        ("围护体系", "地下连续墙及内支撑" if project.retaining_system else "未生成"),
    ])
    document.add_paragraph()
    conclusion_p = document.add_paragraph()
    conclusion_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conclusion_p.add_run(_governing_conclusion(project, checks))
    run.bold = True
    run.font.size = Pt(13)
    document.add_paragraph(DISCLAIMER)
    document.add_page_break()

    # 1. Executive summary
    document.add_heading("1 项目管理摘要", level=1)
    document.add_paragraph(_governing_conclusion(project, checks))
    gv = latest.governing_values if latest else None
    controlling_stability = stability.get("summary", {}).get("controllingLabel")
    _add_key_value_table(document, [
        ("围护墙数量", len(project.retaining_system.diaphragm_walls) if project.retaining_system else 0),
        ("支撑 / 立柱", f"{len(project.retaining_system.supports) if project.retaining_system else 0} / {len(project.retaining_system.columns) if project.retaining_system else 0}"),
        ("最大墙体位移", f"{_text(getattr(gv, 'max_displacement', None))} mm"),
        ("最大墙体弯矩", f"{_text(getattr(gv, 'max_wall_moment', None))} kN·m/m"),
        ("最大支撑轴力", f"{_text(getattr(gv, 'max_support_axial_force', None))} kN"),
        ("稳定性控制项", controlling_stability or "尚未形成"),
        ("检查结果", f"{check_counts.get('fail', 0)} 不通过 / {check_counts.get('warning', 0)} 预警 / {check_counts.get('manual_review', 0)} 人工复核"),
        ("正式出图", "允许" if latest and latest.formal_report_gate and latest.formal_report_gate.allowed_for_official_issue else "暂不允许"),
    ])
    non_pass = [item for item in checks if str(item.get("status")) in {"fail", "warning", "manual_review"}]
    if non_pass:
        document.add_heading("1.1 需要项目负责人关注的事项", level=2)
        grouped_issues = _group_review_checks(non_pass, limit=8)
        _add_table(document, ["等级", "控制问题", "处理建议"], [
            [row[0], f"{row[1]}（控制值/限值：{row[2]}）", f"{row[3]} 由设计人员复核并形成处置记录。"]
            for row in grouped_issues
        ], status_column=0)
    else:
        document.add_paragraph("当前没有记录到不通过或预警项。")

    # 2. Inputs
    document.add_heading("2 工程条件与输入资料", level=1)
    excavation = project.excavation
    geology_audit = project.geological_model.coverage_audit if project.geological_model else {}
    _add_key_value_table(document, [
        ("钻孔数量", len(project.boreholes)),
        ("统一地层数量", len(project.strata)),
        ("轮廓点数量", len(excavation.outline.points) if excavation else 0),
        ("坑顶 / 坑底标高", f"{_text(getattr(excavation, 'top_elevation', None))} / {_text(getattr(excavation, 'bottom_elevation', None))} m"),
        ("地下水位", f"{_text(project.design_settings.groundwater_level)} m"),
        ("地面超载", f"{_text(project.design_settings.surcharge)} kPa"),
        ("地质模型覆盖", "通过" if geology_audit.get("designDomainCovered") else _status(geology_audit.get("status"))),
        ("最大外推距离", f"{_text(geology_audit.get('maximumExtrapolationDistanceM'))} m"),
    ])
    if project.geological_model and project.geological_model.warnings:
        document.add_paragraph("地质模型提示：" + "；".join(project.geological_model.warnings[:5]))

    # 3. Standards guidance
    document.add_heading("3 设计依据与规范对应", level=1)
    document.add_heading("3.1 已确认设计基准", level=2)
    summary_basis = design_basis.get("summary") or {}
    _add_key_value_table(document, [
        ("工程等级 / 基坑安全等级", f"{design_basis.get('projectGrade')} / {design_basis.get('excavationSafetyLevel')}"),
        ("场地复杂程度 / 周边环境", f"{design_basis.get('siteComplexity')} / {design_basis.get('surroundingEnvironmentLevel')}"),
        ("设计阶段", design_basis.get("designStageLabel")),
        ("规范体系", design_basis.get("standardProfileLabel")),
        ("荷载组合策略", design_basis.get("loadCombinationPolicyLabel")),
        ("分项与组合系数", f"γG={_text(summary_basis.get('gammaG'))}, γQ={_text(summary_basis.get('gammaQ'))}, ψ={_text(summary_basis.get('psi'))}"),
        ("重要性系数 / 安全附加储备", f"{_text(summary_basis.get('importanceFactor'))} / {_text(float(summary_basis.get('stabilityReserveRatio') or 0) * 100)}%"),
        ("材料与保护层", f"{summary_basis.get('concreteGrade')} / {summary_basis.get('rebarGrade')} / {summary_basis.get('coverMm')} mm"),
        ("企业工程资源库", f"{summary_basis.get('enterpriseLibraryId') or '-'} / {summary_basis.get('localStandardTemplateId') or '-'}"),
    ])
    document.add_heading("3.2 荷载组合", level=2)
    _add_table(document, ["组合", "表达式", "γG", "γQ", "ψ", "用途"], [
        [row.get("name"), row.get("expression"), row.get("gammaG"), row.get("gammaQ"), row.get("psi"), row.get("note") or "按当前设计基准形成作用效应"]
        for row in design_basis.get("loadCombinations") or []
    ], font_size=8.0)
    document.add_heading("3.3 核心规范对应", level=2)
    document.add_paragraph("下表只列出本项目核心流程中直接参与设计判断的规范。完整条文、地方标准和项目专项要求仍应由设计人员确认。")
    standard_rows: list[list[Any]] = []
    stage_labels = {"input": "工程输入", "scheme": "围护方案", "calculation": "计算验算", "reinforcement": "配筋深化", "deliverables": "成果交付"}
    for stage_key, refs in standards.items():
        # Main report keeps only the two most decision-relevant references per
        # stage; the complete process matrix is retained in Appendix A.
        for ref in refs[:2]:
            standard_rows.append([stage_labels.get(stage_key, stage_key), ref.get("code"), ref.get("name"), ref.get("focus"), ref.get("levelLabel")])
    _add_table(document, ["设计步骤", "规范编号", "规范名称", "本步骤关注内容", "属性"], standard_rows, font_size=7.8)

    document.add_heading("3.4 基坑工程完整验算矩阵", level=2)
    verification_rows = list(verification.get("records") or [])
    _add_table(document, ["类别", "校核项", "设计值/限值", "安全系数", "证据", "状态", "缺口与补齐动作", "规范"], [
        [
            {"strength": "强度", "stiffness": "刚度", "stability": "稳定性", "hydraulic": "水控制", "constructability": "施工性", "other": "其他"}.get(row.get("category"), row.get("category")),
            row.get("label"),
            f"{row.get('designValue') if row.get('designValue') is not None else '—'} / {row.get('limitValue') if row.get('limitValue') is not None else '—'}",
            row.get("safetyFactor") if row.get("safetyFactor") is not None else "—",
            _status(row.get("evidenceState")), _status(row.get("status")),
            row.get("nextAction") or row.get("message") or "—", row.get("standard") or "—",
        ] for row in verification_rows
    ], font_size=6.7, status_column=5)
    missing_input_rows = list(verification.get("missingInputSummary") or [])
    if missing_input_rows:
        document.add_heading("3.5 缺资料闭合清单", level=2)
        _add_table(document, ["资料", "提供阶段", "责任方", "设计阶段可提供", "影响校核", "补齐动作"], [
            [
                row.get("label"), row.get("stageLabel"), row.get("provider"),
                "是" if row.get("designStageAvailable") else "否",
                row.get("affectedCheckCount"), row.get("action"),
            ] for row in missing_input_rows
        ], font_size=7.0)
    wall_verification_rows = list(verification.get("wallObjects") or [])
    if wall_verification_rows:
        document.add_heading("3.6 逐墙验算证据摘要", level=2)
        _add_table(document, ["墙对象", "类型", "墙厚(m)", "顶/底标高(m)", "已计算", "不通过", "待闭合", "状态"], [
            [
                row.get("wallCode"), row.get("wallTypeLabel"), row.get("thicknessM"),
                f"{row.get('topElevationM')} / {row.get('bottomElevationM')}",
                (row.get("summary") or {}).get("calculatedCount"),
                (row.get("summary") or {}).get("failCount"),
                (row.get("summary") or {}).get("reviewCount"), _status(row.get("status")),
            ] for row in wall_verification_rows
        ], font_size=7.2, status_column=7)

    # 4. Schemes
    document.add_heading("4 围护结构多方案比选", level=1)
    rows = list(scheme_comparison.get("rows") or [])
    if rows:
        document.add_paragraph("候选方案先通过平面拓扑、围檩支点、障碍避让和节点传力预检。完整比选完成后，再依据位移、轴力、围檩内力和稳定性确定推荐方案。")
        _add_table(document, ["方案", "体系", "支撑/立柱", "最长跨度(m)", "最大轴力(kN)", "最大位移(mm)", "最小稳定系数", "排名"], [
            [
                f"方案 {row.get('schemeLabel')}{'（推荐）' if row.get('recommended') else ''}",
                row.get("schemeName"),
                f"{row.get('supportCount')} / {row.get('columnCount')}",
                row.get("maxSpanLength"),
                row.get("maxSupportAxialForce") if row.get("fullCalculationReady") else "待完整计算",
                row.get("maxDisplacement") if row.get("fullCalculationReady") else "待完整计算",
                row.get("minStabilitySafetyFactor") if row.get("fullCalculationReady") else "待完整计算",
                row.get("decisionRank") or row.get("rank"),
            ] for row in rows
        ], font_size=7.5)
        _add_picture(document, chart_by_title.get("支撑优化候选方案平面比选图"), "图 4-1 A/B/C 支撑方案平面比选")
        _add_picture(document, chart_by_title.get("支撑优化候选方案评分图"), "图 4-2 候选方案拓扑预检评分")
    else:
        document.add_paragraph("尚未生成可比选的围护方案。")

    # 5. Adopted design
    document.add_heading("5 当前采用围护结构", level=1)
    retaining = project.retaining_system
    if retaining:
        _add_key_value_table(document, [
            ("围护墙", len(retaining.diaphragm_walls)),
            ("冠梁", len(retaining.crown_beams)),
            ("围檩", len(retaining.wale_beams)),
            ("水平支撑", len(retaining.supports)),
            ("临时立柱", len(retaining.columns)),
            ("支撑节点", len(retaining.support_nodes)),
        ])
        wall_rows = [[wall.panel_code, wall.design_face_code or wall.segment_id, wall.thickness, wall.top_elevation, wall.bottom_elevation, wall.concrete_grade, _status(getattr(wall.design_results, "check_status", None))] for wall in retaining.diaphragm_walls[:30]]
        _add_table(document, ["墙段", "边段", "厚度(m)", "墙顶标高", "墙底标高", "混凝土", "设计状态"], wall_rows, font_size=7.8)
        _add_picture(document, chart_by_title.get("支撑布置评分平面图"), "图 5-1 当前采用方案支撑平面")
    else:
        document.add_paragraph("尚未生成围护结构。")

    # 6. Calculation results
    document.add_heading("6 施工阶段内力与变形", level=1)
    if latest:
        _add_key_value_table(document, [
            ("最大合成侧压力", f"{_text(gv.max_total_pressure)} kPa"),
            ("最大墙体位移", f"{_text(gv.max_displacement)} mm"),
            ("最大墙体弯矩", f"{_text(gv.max_wall_moment)} kN·m/m"),
            ("最大墙体剪力", f"{_text(gv.max_wall_shear)} kN/m"),
            ("最大支撑轴力", f"{_text(gv.max_support_axial_force)} kN"),
            ("控制检查状态", _status(gv.governing_check_status)),
        ])
        case = next((item for item in project.calculation_cases if item.id == latest.case_id), project.calculation_cases[-1] if project.calculation_cases else None)
        stage_label_by_id: dict[str, str] = {}
        if case:
            for index, stage_def in enumerate(case.stages):
                name_lower = str(getattr(stage_def, "name", "") or "").lower()
                if "final" in name_lower or "service" in name_lower or index == len(case.stages) - 1:
                    label = "最终开挖与使用工况"
                else:
                    elevation = getattr(stage_def, "excavation_elevation", None)
                    label = f"第 {index + 1} 阶段（开挖至 {_text(elevation)} m）"
                stage_label_by_id[str(stage_def.id)] = label
        stage_rows = []
        for stage in latest.stage_results[:30]:
            max_pressure = max((point.total_pressure for point in stage.pressure_profile.points), default=0.0)
            wall = stage.wall_internal_force
            stage_rows.append([
                stage_label_by_id.get(str(stage.stage_id), str(stage.stage_id)),
                stage.segment_id,
                max_pressure,
                wall.max_moment if wall else None,
                wall.max_shear if wall else None,
                wall.max_displacement if wall else None,
                len(stage.support_forces),
            ])
        _add_table(document, ["施工阶段", "边段", "最大压力(kPa)", "墙弯矩", "墙剪力", "墙位移(mm)", "支撑结果数"], stage_rows, font_size=7.4)
        for title, caption in (
            ("墙体土压力图", "图 6-1 控制阶段墙体侧压力分布"),
            ("墙体位移", "图 6-2 控制阶段墙体位移包络"),
            ("墙体弯矩", "图 6-3 控制阶段墙体弯矩包络"),
            ("墙体剪力", "图 6-4 控制阶段墙体剪力包络"),
            ("支撑轴力柱状图", "图 6-5 支撑设计轴力分布"),
            ("围檩弯矩包络图", "图 6-6 围檩弯矩包络"),
        ):
            _add_picture(document, chart_by_title.get(title), caption)
    else:
        document.add_paragraph("尚未运行施工阶段计算。")

    # 7. Stability
    document.add_heading("7 稳定与水控制完整验算", level=1)
    factors = list(stability.get("factors") or [])
    if factors:
        summary = stability.get("summary") or {}
        document.add_paragraph(
            f"共列出 {summary.get('count', 0)} 项稳定与水控制项目，其中已计算 {summary.get('calculatedCount', 0)} 项、待补资料或计算 {summary.get('pendingCount', 0)} 项。控制项为“{summary.get('controllingLabel') or '-'}”，"
            f"最小限值比为 { _text(summary.get('minimumMarginRatio')) }。限值比小于 1.0 表示不满足，"
            "1.0～1.10 作为优先复核区间。"
        )
        _add_table(document, ["验算项目", "计算系数", "规范限值", "限值比", "证据/状态", "缺口与动作", "规范依据"], [
            [item.get("label"), item.get("value"), item.get("limit"), item.get("marginRatio"), _status(item.get("evidenceState") or item.get("status")), item.get("nextAction") or item.get("message") or "—", f"{item.get('standard')}；{item.get('clauseFocus')}"]
            for item in factors
        ], font_size=7.0, status_column=4)
        _add_picture(document, chart_by_title.get("稳定性安全系数分布图"), "图 7-1 稳定性安全系数整体分布")
    else:
        document.add_paragraph("尚未形成可用的稳定性检查结果。")

    document.add_heading("7.2 正式不利工况专项复算", level=2)
    scenario_rows = list(formal_scenarios.get("summaries") or [])
    if scenario_rows:
        _add_table(document, ["不利工况", "最大墙位移(mm)", "最大支撑轴力(kN)", "最小安全系数", "状态", "证据"], [
            [row.get("scenarioLabel"), row.get("maxWallDisplacementMm"), row.get("maxSupportForceKn"), row.get("minimumSafetyFactor"), _status(row.get("status")), "独立施工阶段正式复算"]
            for row in scenario_rows
        ], font_size=7.5, status_column=4)
        errors = list(formal_scenarios.get("errors") or [])
        if errors:
            document.add_paragraph("未完成场景：" + "；".join(f"{row.get('scenarioLabel') or row.get('scenarioCode')}：{_short(row.get('error'), 80)}" for row in errors[:6]))
    else:
        document.add_paragraph("尚未执行正式不利工况专项复算；当前稳定性结论仅对应已计算的基准工况。")

    # 8. Reinforcement
    document.add_heading("8 配筋设计与构造摘要", level=1)
    rebar_scheme = retaining.rebar_design_scheme if retaining else {}
    if rebar_scheme:
        _add_key_value_table(document, [
            ("配筋模式", rebar_scheme.get("mode") or "balanced"),
            ("总体状态", _status(rebar_scheme.get("status"))),
            ("墙体配筋分区", len(rebar_scheme.get("wallZones") or [])),
            ("支撑配筋组", len(rebar_scheme.get("supportSchemes") or [])),
            ("节点附加筋组", len(rebar_scheme.get("beamNodeSchemes") or [])),
            ("配筋检查数量", len(rebar_scheme.get("checks") or [])),
        ])
        wall_rebar_rows: list[list[Any]] = []
        for wall in retaining.diaphragm_walls[:40]:
            for group in wall.reinforcement[:8]:
                wall_rebar_rows.append([wall.panel_code, group.name, group.grade, group.diameter, group.spacing or group.count, group.required_area_per_meter, group.area_per_meter, group.check_status])
        _add_table(document, ["构件", "钢筋组", "等级", "直径(mm)", "间距/根数", "需求面积", "提供面积", "状态"], wall_rebar_rows, font_size=7.6, status_column=7)
        wale_rows = []
        for beam in retaining.wale_beams:
            result = beam.design_result
            if result:
                wale_rows.append([beam.code, beam.support_level, result.max_moment_design, result.max_shear_design, result.required_reinforcement_area, result.provided_reinforcement_area, f"Φ{result.main_bar_diameter}@{result.main_bar_spacing}", result.check_status])
        _add_table(document, ["围檩", "层号", "设计弯矩", "设计剪力", "As需求", "As提供", "主筋", "状态"], wale_rows, font_size=7.6, status_column=7)
        limitations = list(rebar_scheme.get("limitations") or [])
        if limitations:
            document.add_paragraph("配筋深化边界：" + "；".join(_short(item, 120) for item in limitations[:6]))
        deepening_gate = dict((rebar_scheme.get("diagnostics") or {}).get("deepeningGate") or {})
        if deepening_gate:
            document.add_heading("8.2 配筋深化入口诊断", level=2)
            _add_key_value_table(document, [
                ("入口状态", deepening_gate.get("status")),
                ("入口总阻断", deepening_gate.get("blockerCount")),
                ("发行阻断", deepening_gate.get("releaseBlockerCount")),
                ("复核项", deepening_gate.get("warningCount")),
                ("可进入 P3", "是" if deepening_gate.get("canRunP3") else "否"),
                ("可发行施工图", "是" if deepening_gate.get("canIssueConstructionDrawings") else "否"),
            ])
            gate_rows = list(deepening_gate.get("blockers") or []) + list(deepening_gate.get("warnings") or [])
            if gate_rows:
                _add_table(document, ["原因", "数量", "影响对象", "说明", "补齐动作", "目标阶段"], [
                    [
                        row.get("title"), row.get("count"), "、".join(row.get("objects") or []) or "—",
                        row.get("message"), row.get("requiredAction"), row.get("targetStage"),
                    ] for row in gate_rows[:24]
                ], font_size=7.0)
        document.add_heading("8.3 企业节点、预埋件与钢筋空间深化" if deepening_gate else "8.2 企业节点、预埋件与钢筋空间深化", level=2)
        if p3_closure:
            p3_summary = dict(p3_closure.get("summary") or {})
            _add_key_value_table(document, [
                ("深化状态", _status(p3_closure.get("status"))),
                ("逐根钢筋数量", p3_summary.get("individualBarCount")),
                ("机械连接数量", p3_summary.get("couplerCount")),
                ("预埋件数量", p3_summary.get("embeddedItemCount")),
                ("高风险节点局部模型", p3_summary.get("nodeSubmodelCount")),
                ("硬碰撞 / 协调问题", f"{p3_summary.get('hardCollisionCount', 0)} / {p3_summary.get('coordinationIssueCount', 0)}"),
                ("企业节点模板未覆盖", p3_summary.get("unmatchedEnterpriseNodeCount")),
            ])
            checks_p3 = list(p3_closure.get("controllingChecks") or [])
            if checks_p3:
                _add_table(document, ["构件", "类别", "计算值", "限值", "状态", "处理建议"], [
                    [row.get("hostCode") or row.get("hostId"), row.get("category"), row.get("calculatedValue"), row.get("limitValue"), _status(row.get("status")), row.get("recommendedAction") or row.get("message")]
                    for row in checks_p3[:20]
                ], font_size=7.2, status_column=4)
        else:
            document.add_paragraph("尚未执行企业节点、预埋件与钢筋空间深化闭环。")
    else:
        document.add_paragraph("尚未生成配筋方案。")

    # 9. Review list
    document.add_heading("9 校核结论与复核清单", level=1)
    document.add_paragraph(_governing_conclusion(project, checks))
    document.add_paragraph("重复检查已按规则、状态和结论合并，表中数值为该组控制值；完整逐构件记录保留在审计归档中。")
    _add_table(
        document,
        ["状态", "检查项目", "控制值 / 限值", "问题及影响范围", "规范依据"],
        _group_review_checks(checks, limit=8),
        font_size=7.6,
        status_column=0,
    )

    # 10. Deliverables
    document.add_heading("10 成果使用与责任边界", level=1)
    document.add_paragraph(
        "项目管理者应优先审阅本计算书第1、4、6、7、8和9章。施工图、BIM模型、钢筋表和机器可读归档数据"
        "属于同一设计快照的不同表达形式。任何围护几何、支撑体系、地层参数、地下水位或施工阶段变更后，"
        "原计算结果和配筋结论均应重新校核。"
    )
    _add_table(document, ["成果", "主要使用者", "用途", "注意事项"], [
        ["本计算书", "项目负责人、设计、校核", "方案决策、控制结果和风险复核", "以明文结论和图表为主"],
        ["施工图", "设计、施工、监理", "构件定位、尺寸、配筋和节点表达", "须完成校审和发行签署"],
        ["IFC/BIM", "设计协调、施工技术", "空间协调和模型交底", "不得单独替代施工图"],
        ["机器归档数据", "系统管理员、二次开发", "追溯、迁移和复算", "不作为项目管理者主阅读文件"],
    ], font_size=8)

    # Required technical appendices kept compact and separated from main management narrative.
    document.add_page_break()
    document.add_heading("附录 A 设计流程—关键计算—规范条文对应矩阵", level=1)
    matrix = build_standards_process_matrix(project)
    matrix_rows = []
    for step in matrix.get("steps") or []:
        matrix_rows.append([
            step.get("index"),
            step.get("title"),
            "；".join(step.get("keyCalculations") or []),
            "；".join(ref.get("code") for ref in step.get("standardRefs") or []),
            "；".join(step.get("outputs") or []),
        ])
    _add_table(document, ["序号", "流程", "关键计算", "主要规范", "输出"], matrix_rows, font_size=7.0)

    document.add_heading("附录 B 关键计算原理、公式与复核点", level=1)
    _add_table(document, ["计算项目", "核心表达", "重点复核"], [
        ["土压力与水压力", "分层有效应力积分，土水压力按工况组合", "土参数代表性、地下水位和地面超载"],
        ["墙体内力与变形", "墙体梁-土弹簧-支撑弹簧施工阶段求解", "支撑激活顺序、土弹簧、位移控制值"],
        ["围檩与支撑", "围檩连续梁反力与支撑轴力协调", "节点传力、有效长度、预加轴力和温度效应"],
        ["稳定性", "嵌固、隆起、渗流、突涌和整体稳定分别验算", "控制剖面、不利土层和安全系数限值"],
        ["混凝土配筋", "按内力包络计算需求钢筋并检查最小配筋与构造", "裂缝、锚固、搭接、节点和施工可实施性"],
    ], font_size=7.6)

    document.add_heading("附录 C 计算追溯信息", level=1)
    manifest = version_manifest()
    _add_key_value_table(document, [
        ("软件版本", manifest.get("softwareVersion")),
        ("算法版本", manifest.get("algorithmVersion")),
        ("规则库版本", manifest.get("ruleSetVersion")),
        ("导出协议版本", manifest.get("exportSchemaVersion")),
        ("输入快照哈希", getattr(latest, "input_snapshot_hash", None) if latest else None),
        ("设计快照哈希", getattr(latest, "adopted_design_snapshot_hash", None) if latest else None),
        ("结果哈希", getattr(latest, "result_hash", None) if latest else None),
        ("计算合同", getattr(latest, "calculation_contract_id", None) if latest else None),
    ])

    document.save(path)
    return path
